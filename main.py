import numpy as np
import docx
from docx.shared import Pt, Mm
import os

local_path = "/work/path"
freq_file = local_path + "test_files/structure-11605666-pm3-dft1-b3lyp-631++gdp-freq.log"
geo_file = local_path + "test_files/structure-11605666-pm3-dft1-b3lyp-631++gdp.log"

# Delimiting Parameters
sep_arg_one = "normal coordinates:"
sep_arg_two = "Thermochemistry"
sep_arg_three = "Optimized Parameters"
sep_arg_four = "Stoichiometry"

freq_arr = []
ir_arr = []
geo_r_arr = []
geo_a_arr = []
geo_d_arr = []


# Dual and Uniform scaling multiplier values

# For Dual 1 -> value > 1800 cm^-1 [0.9659] & For Dual 1 -> value < 1800 cm^-1 [0.9927]

def get_dual_one_multipllier(value):
    if value > 1800:
        return round(value * 0.9659)
    else:
        return round(value * 0.9927)


# For Dual 2 -> value > 1800 cm^-1 [0.955] & For Dual 2 -> value < 1800 cm^-1 [0.977]

def get_dual_two_multiplier(value):
    if value > 1800:
        return round(value * 0.955)
    else:
        return round(value * 0.977)


# For Uniform

def get_uniform_multiplier(value):
    return round(value * 0.9726)


# It performs the process of creating a numpy array with an index containing three parameters.

def connect_list(arr):
    new_list = []
    for i in range(0, len(arr)):
        sub_list = []
        for j in range(0, len(arr), round(len(arr) / 3) + 1):
            if i + j < len(arr):
                sub_list.extend(arr[i + j])
            else:
                sub_list.append(" ")
                sub_list.append(" ")
                sub_list.append(" ")

        new_list.append(sub_list)

    return new_list[:(round(len(new_list) / 3) + 1)]


# "!" in the list deletes the sign and spaces

def edit_list(geo_arr):
    arr = []

    for i in geo_arr:
        str_geo = i.replace("!", "").split(" ")
        while "" in str_geo:
            str_geo.remove("")

        arr.append(str_geo[:3])

    final_list_geo = connect_list(arr)
    final_list_geo = tuple([tuple(row) for row in final_list_geo])
    return final_list_geo

os.chdir(local_path)


def create_freq_table(data):
    doc = docx.Document()
    doc.add_heading('Optimized Frequency Table', 0)

    # Table data in a form of list
    data = data

    # Creating a table object
    table = doc.add_table(rows=1, cols=7)

    # Adding heading in the 1st row of the table
    row = table.rows[0].cells

    header = ["ID", "File_ID", "Frequency", "Dual S^1", "Uniform", "IR Int.", "Dual S^2"]
    header_var = ["zero", "zero_t", "one", "two", "three", "four", "five"]

    for i in range(0, len(header)):
        header_var[i] = row[i].paragraphs[0].add_run(str(header[i]))
        header_var[i].bold = True

    # Adding data from the list to the table
    for id, file_id, frequency, dual_1, uniform, ir, dual_2 in data:
        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(id)
        row[1].text = str(file_id)
        row[2].text = str(frequency)
        row[3].text = str(dual_1)
        row[4].text = str(uniform)
        row[5].text = str(ir)
        row[6].text = str(dual_2)

    table.style = 'Table Grid'
    doc.save('optimized_freq_table.docx')


def create_geo_table(data_r, data_a, data_d):
    doc = docx.Document()
    doc.add_heading('Optimized Geometry Table', 0)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)

    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)

    # Table data in a form of list
    data_r = data_r
    data_a = data_a
    data_d = data_d

    # Creating a table object
    table_one = doc.add_table(rows=1, cols=9, style='Table Grid')
    doc.add_page_break()
    table_two = doc.add_table(rows=1, cols=9, style='Table Grid')
    doc.add_page_break()
    table_three = doc.add_table(rows=1, cols=9, style='Table Grid')

    # Adding heading in the 1st row of the table
    row_r = table_one.rows[0].cells
    row_a = table_two.rows[0].cells
    row_d = table_three.rows[0].cells

    # Col merging for R
    table_r_a, table_r_b = row_r[:2]
    table_r_a.merge(table_r_b)
    table_r_c, table_r_d = row_r[3:5]
    table_r_c.merge(table_r_d)
    table_r_e, table_r_f = row_r[6:8]
    table_r_e.merge(table_r_f)

    # Col merging for A
    table_a_a, table_a_b = row_a[:2]
    table_a_a.merge(table_a_b)
    table_a_c, table_a_d = row_a[3:5]
    table_a_c.merge(table_a_d)
    table_a_e, table_a_f = row_a[6:8]
    table_a_e.merge(table_a_f)

    # Col merging for D
    table_d_a, table_d_b = row_d[:2]
    table_d_a.merge(table_d_b)
    table_d_c, table_d_d = row_d[3:5]
    table_d_c.merge(table_d_d)
    table_d_e, table_d_f = row_d[6:8]
    table_d_e.merge(table_d_f)

    # Adding heading in the 1st row of the table for D
    row_d[0].text = 'Dihedral'
    row_d[2].text = 'Dihedral Dönü Derecesi'
    row_d[3].text = 'Dihedral'
    row_d[5].text = 'Dihedral Dönü Derecesi'
    row_d[6].text = 'Dihedral'
    row_d[8].text = 'Dihedral Dönü Derecesi'

    # Adding data from the list to the table for R
    for aa, bb, cc, dd, ee, ff, gg, hh, jj in data_r:
        # Adding a row and then adding data in it.
        row = table_one.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(aa)
        row[1].text = str(bb)
        row[2].text = str(cc)
        row[3].text = str(dd)
        row[4].text = str(ee)
        row[5].text = str(ff)
        row[6].text = str(gg)
        row[7].text = str(hh)
        row[8].text = str(jj)

    # Adding data from the list to the table for A
    for aa, bb, cc, dd, ee, ff, gg, hh, jj in data_a:
        # Adding a row and then adding data in it.
        row = table_two.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(aa)
        row[1].text = str(bb)
        row[2].text = str(cc)
        row[3].text = str(dd)
        row[4].text = str(ee)
        row[5].text = str(ff)
        row[6].text = str(gg)
        row[7].text = str(hh)
        row[8].text = str(jj)

    # Adding data from the list to the table for D
    for aa, bb, cc, dd, ee, ff, gg, hh, jj in data_d:
        # Adding a row and then adding data in it.
        row = table_three.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(aa)
        row[1].text = str(bb)
        row[2].text = str(cc)
        row[3].text = str(dd)
        row[4].text = str(ee)
        row[5].text = str(ff)
        row[6].text = str(gg)
        row[7].text = str(hh)
        row[8].text = str(jj)

    doc.save('optimized_geo_table.docx')


def get_freq_data():
    with open(freq_file, 'r', encoding='utf-8') as file:
        contents = file.read()
        result = contents.split(sep_arg_one)[1]

    new_result = result.split(sep_arg_two)[0]

    for i in new_result.splitlines():
        if "Frequencies" in i:
            freq_arr.append(i)
        if "IR Inten" in i:
            ir_arr.append(i)

    arr = []
    for i in freq_arr:
        str_freq = i.replace("Frequencies --", "").split(" ")
        while ("" in str_freq):
            str_freq.remove("")
        arr.append(str_freq)

    last_freq_arr = np.around(np.array(arr).flatten().astype(np.float64), 0)
    last_freq_arr = np.flip(np.array(last_freq_arr, dtype=np.int16))

    arr = []
    for i in ir_arr:
        str_ir = i.replace("IR Inten    --", "").split(" ")
        while ("" in str_ir):
            str_ir.remove("")
        arr.append(str_ir)

    last_ir_arr = np.around(np.array(arr).flatten().astype(np.float64), 0)
    last_ir_arr = np.flip(np.array(last_ir_arr, dtype=np.int16))

    new_data = []
    count = len(last_freq_arr)
    for i in range(0, len(last_freq_arr)):
        new_data.append([i + 1, count, last_freq_arr[i], get_dual_one_multipllier(last_freq_arr[i]),
                         get_uniform_multiplier(last_freq_arr[i]), last_ir_arr[i],
                         get_dual_two_multiplier(last_freq_arr[i])])

        count = count - 1

    result = tuple([tuple(row) for row in new_data])
    create_freq_table(result)


def get_geo_file():
    with open(geo_file, 'r', encoding='utf-8') as file:
        contents = file.read()
        result = contents.split(sep_arg_three)[1]

    new_result = result.split(sep_arg_four)[0]

    for i in new_result.splitlines()[5:-4]:

        if "R" in i:
            geo_r_arr.append(i)
        if "A" in i:
            geo_a_arr.append(i)
        if " D" in i:
            geo_d_arr.append(i)

    result_r = edit_list(geo_r_arr)
    result_a = edit_list(geo_a_arr)
    result_d = edit_list(geo_d_arr)

    create_geo_table(result_r, result_a, result_d)


try:
    get_freq_data()
    get_geo_file()

    print("A word document was created for the frequency table.")
except Exception as e:
    print(e)